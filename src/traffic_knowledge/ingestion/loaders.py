"""Safe text extraction for supported traffic knowledge documents."""

from __future__ import annotations

import re
from collections.abc import Callable
from pathlib import Path

import pymupdf
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from traffic_knowledge.domain.document import (
    DocumentValidationError,
    ParsedDocument,
    SourceBlock,
)

MARKDOWN_MEDIA_TYPE = "text/markdown"
DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
PDF_MEDIA_TYPE = "application/pdf"
_MARKDOWN_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_DOCX_HEADING = re.compile(r"^Heading\s+([1-6])$", re.IGNORECASE)

Loader = Callable[[Path], ParsedDocument]


def _normalize_text(text: str) -> str:
    return " ".join(text.split())


def _heading_location(headings: list[str], fallback: str = "document") -> str:
    active = [heading for heading in headings if heading]
    return " > ".join(active) if active else fallback


def _set_heading(headings: list[str], level: int, title: str) -> None:
    del headings[level - 1 :]
    while len(headings) < level - 1:
        headings.append("")
    headings.append(_normalize_text(title))


def _load_markdown(path: Path) -> ParsedDocument:
    try:
        lines = path.read_text(encoding="utf-8").splitlines()
    except (OSError, UnicodeError) as error:
        raise DocumentValidationError("DOCUMENT_UNREADABLE", path.name) from error

    headings: list[str] = []
    section_lines: list[str] = []
    blocks: list[SourceBlock] = []

    def flush_section() -> None:
        text = _normalize_text("\n".join(section_lines))
        if text:
            blocks.append(
                SourceBlock(
                    text=text,
                    location=_heading_location(headings),
                    ordinal=len(blocks),
                )
            )
        section_lines.clear()

    for line in lines:
        match = _MARKDOWN_HEADING.match(line)
        if match:
            flush_section()
            _set_heading(headings, len(match.group(1)), match.group(2))
        else:
            section_lines.append(line)
    flush_section()

    return ParsedDocument(path.name, MARKDOWN_MEDIA_TYPE, tuple(blocks))


def _load_docx(path: Path) -> ParsedDocument:
    try:
        document = Document(path)
    except (OSError, ValueError, KeyError, PackageNotFoundError) as error:
        raise DocumentValidationError("DOCUMENT_DOCX_UNREADABLE", path.name) from error

    headings: list[str] = []
    section_lines: list[str] = []
    blocks: list[SourceBlock] = []

    def flush_section() -> None:
        text = _normalize_text("\n".join(section_lines))
        if text:
            blocks.append(
                SourceBlock(
                    text=text,
                    location=_heading_location(headings),
                    ordinal=len(blocks),
                )
            )
        section_lines.clear()

    for paragraph in document.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        heading_match = _DOCX_HEADING.match(style_name)
        if heading_match:
            flush_section()
            _set_heading(headings, int(heading_match.group(1)), paragraph.text)
        else:
            section_lines.append(paragraph.text)
    flush_section()

    return ParsedDocument(path.name, DOCX_MEDIA_TYPE, tuple(blocks))


def _load_pdf(path: Path) -> ParsedDocument:
    try:
        document = pymupdf.open(path)
    except (OSError, RuntimeError, ValueError, pymupdf.FileDataError) as error:
        raise DocumentValidationError("DOCUMENT_PDF_UNREADABLE", path.name) from error

    try:
        if document.needs_pass:
            raise DocumentValidationError("DOCUMENT_PDF_ENCRYPTED", path.name)

        blocks = []
        for page_number, page in enumerate(document, start=1):
            text = _normalize_text(page.get_text("text"))
            if text:
                blocks.append(
                    SourceBlock(
                        text=text,
                        location=f"page:{page_number}",
                        ordinal=len(blocks),
                    )
                )
    except DocumentValidationError:
        raise
    except (OSError, RuntimeError, ValueError) as error:
        raise DocumentValidationError("DOCUMENT_PDF_UNREADABLE", path.name) from error
    finally:
        document.close()

    return ParsedDocument(path.name, PDF_MEDIA_TYPE, tuple(blocks))


LOADERS: dict[str, Loader] = {
    ".md": _load_markdown,
    ".docx": _load_docx,
    ".pdf": _load_pdf,
}


def load_document(path: Path) -> ParsedDocument:
    """Parse one supported file without executing embedded content."""
    path = Path(path)
    if not path.is_file():
        raise DocumentValidationError("DOCUMENT_NOT_FOUND", str(path))

    loader = LOADERS.get(path.suffix.lower())
    if loader is None:
        raise DocumentValidationError("DOCUMENT_TYPE_UNSUPPORTED", path.suffix)

    parsed = loader(path)
    if not any(block.text.strip() for block in parsed.blocks):
        raise DocumentValidationError("DOCUMENT_EMPTY", path.name)
    return parsed
