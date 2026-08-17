from pathlib import Path

import pymupdf
import pytest
from docx import Document

from traffic_knowledge.domain.document import DocumentValidationError
from traffic_knowledge.ingestion.loaders import load_document


def test_loads_reviewable_markdown_fixture():
    path = Path(__file__).parent / "fixtures" / "sample.md"

    parsed = load_document(path)

    assert parsed.blocks[0].location == "Traffic Operations > Forecasting Metrics"
    assert parsed.blocks[0].text == "MAE measures the average absolute prediction error."


def test_loads_markdown_with_nested_heading_location(tmp_path):
    path = tmp_path / "guide.md"
    path.write_text(
        "# Traffic\n\nOverview.\n\n## GRU\nShort-term   model.\n",
        encoding="utf-8",
    )

    parsed = load_document(path)

    assert parsed.filename == "guide.md"
    assert parsed.media_type == "text/markdown"
    assert [block.location for block in parsed.blocks] == ["Traffic", "Traffic > GRU"]
    assert parsed.blocks[1].text == "Short-term model."


def test_loads_docx_with_heading_location(tmp_path):
    path = tmp_path / "guide.docx"
    document = Document()
    document.add_heading("Traffic", level=1)
    document.add_heading("Metrics", level=2)
    document.add_paragraph("MAE measures absolute prediction error.")
    document.save(path)

    parsed = load_document(path)

    assert parsed.media_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(parsed.blocks) == 1
    assert parsed.blocks[0].location == "Traffic > Metrics"
    assert parsed.blocks[0].text == "MAE measures absolute prediction error."


def test_loads_pdf_with_page_location(tmp_path):
    path = tmp_path / "guide.pdf"
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "PEMS04 contains traffic flow observations.")
    document.save(path)
    document.close()

    parsed = load_document(path)

    assert parsed.media_type == "application/pdf"
    assert len(parsed.blocks) == 1
    assert parsed.blocks[0].location == "page:1"
    assert "PEMS04 contains traffic flow observations." in parsed.blocks[0].text


def test_rejects_unsupported_extension(tmp_path):
    path = tmp_path / "notes.txt"
    path.write_text("traffic", encoding="utf-8")

    with pytest.raises(DocumentValidationError) as error:
        load_document(path)

    assert error.value.code == "DOCUMENT_TYPE_UNSUPPORTED"


def test_rejects_empty_document(tmp_path):
    path = tmp_path / "empty.md"
    path.write_text("# Heading only\n", encoding="utf-8")

    with pytest.raises(DocumentValidationError) as error:
        load_document(path)

    assert error.value.code == "DOCUMENT_EMPTY"


def test_reports_unreadable_pdf(tmp_path):
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"not a pdf")

    with pytest.raises(DocumentValidationError) as error:
        load_document(path)

    assert error.value.code == "DOCUMENT_PDF_UNREADABLE"


def test_reports_encrypted_pdf(tmp_path):
    path = tmp_path / "encrypted.pdf"
    document = pymupdf.open()
    document.new_page().insert_text((72, 72), "Protected traffic report")
    document.save(
        path,
        encryption=pymupdf.PDF_ENCRYPT_AES_256,
        owner_pw="owner-password",
        user_pw="reader-password",
    )
    document.close()

    with pytest.raises(DocumentValidationError) as error:
        load_document(path)

    assert error.value.code == "DOCUMENT_PDF_ENCRYPTED"


def test_rejects_missing_file(tmp_path):
    path = Path(tmp_path / "missing.md")

    with pytest.raises(DocumentValidationError) as error:
        load_document(path)

    assert error.value.code == "DOCUMENT_NOT_FOUND"
